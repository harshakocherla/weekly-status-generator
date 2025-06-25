#!/usr/bin/env python3

import os
import sys
import git
from datetime import datetime, timedelta
from dateutil.parser import parse
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

def is_monday(date):
    """Check if the given date is a Monday."""
    return date.weekday() == 0

def get_week_dates(monday_date):
    """Get all dates from Monday to Friday for the given week."""
    dates = []
    for i in range(5):  # Monday to Friday
        dates.append(monday_date + timedelta(days=i))
    return dates

def get_weeks_until_today(start_monday):
    """Get all Monday dates from start_monday until today's week."""
    weeks = []
    current_monday = start_monday
    today = datetime.now().date()
    
    while current_monday <= today:
        weeks.append(current_monday)
        current_monday += timedelta(days=7)
    
    return weeks

def is_git_repo(path):
    """Check if the given path is a git repository."""
    try:
        git.Repo(path)
        return True
    except git.InvalidGitRepositoryError:
        return False

def get_current_git_user(repo):
    """Get the current git user's email and name."""
    try:
        email = repo.config_reader().get_value("user", "email", "")
        name = repo.config_reader().get_value("user", "name", "")
        return email, name
    except Exception:
        return None, None

def get_commits_for_date(repo, date, user_email=None, user_name=None):
    """Get all commits for a specific date in the repository."""
    commits = []
    start_date = datetime.combine(date, datetime.min.time())
    end_date = datetime.combine(date, datetime.max.time())
    
    for commit in repo.iter_commits():
        commit_date = datetime.fromtimestamp(commit.committed_date)
        if start_date <= commit_date <= end_date:
            # Filter by user if specified
            if user_email and user_name:
                if (commit.author.email == user_email or 
                    commit.author.name == user_name):
                    commits.append(commit)
            else:
                commits.append(commit)
    
    return commits

def process_workspace(workspace_path, monday_date):
    """Process the workspace and generate status report."""
    workspace = Path(workspace_path)
    if not workspace.exists():
        print(f"Error: Workspace path {workspace_path} does not exist.")
        return

    weeks = get_weeks_until_today(monday_date)
    all_reports = {}

    print("\nScanning repositories in ~/Workspaces:")
    print("-" * 50)
    
    # Process each item in the workspace
    for item in workspace.iterdir():
        if item.is_dir():
            if is_git_repo(item):
                print(f"✓ Found git repository: {item.name}")
                repo = git.Repo(item)
                user_email, user_name = get_current_git_user(repo)
                if user_email or user_name:
                    print(f"  👤 Filtering commits for: {user_name or user_email}")
                
                for week_monday in weeks:
                    week_dates = get_week_dates(week_monday)
                    if week_monday not in all_reports:
                        all_reports[week_monday] = {}
                        for date in week_dates:
                            all_reports[week_monday][date] = "continuing previous day's work"
                    
                    for date in week_dates:
                        commits = get_commits_for_date(repo, date, user_email, user_name)
                        if commits:
                            all_reports[week_monday][date] = commits[0].message.strip()
            else:
                print(f"✗ Not a git repository: {item.name}")
    
    print("-" * 50)
    return all_reports

def generate_report(all_reports, output_file):
    """Generate the status report file."""
    with open(output_file, 'w') as f:
        for week_monday, week_report in sorted(all_reports.items()):
            f.write(f"\nReport for week starting at {week_monday.strftime('%m/%d/%Y')}\n")
            f.write("-" * 50 + "\n")
            for date, message in sorted(week_report.items()):
                f.write(f"{date.strftime('%m/%d/%Y')} - {message}\n")
            f.write("-" * 50 + "\n")
        f.write("\nEND of Report\n")

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in edge_data:
                element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_padding(cell, top=0.08, bottom=0.08, left=0.08, right=0.08):
    """
    Set cell padding in inches (default 0.08in for all sides).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for margin, value in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = tcMar.find(qn('w:' + margin))
        if node is None:
            node = OxmlElement('w:' + margin)
            tcMar.append(node)
        node.set(qn('w:w'), str(int(Inches(value).twips)))
        node.set(qn('w:type'), 'dxa')

def fill_docx_report(template_path, output_path, week_report, week_monday, repo_commits):
    """Fill the docx template with a single week's status report data and save as a new file."""
    doc = Document(template_path)
    # Fill header fields (assume first 2 paragraphs for FROM/WEEK ENDING, next 2 for CLIENT/PROJECT)
    # FROM: Harsha Kocherla
    # WEEK ENDING: last date in week_report
    # CLIENT NAME: Best Buy
    # PROJECT NAME: comma-separated repo names with commits
    week_dates = sorted(week_report.keys())
    week_ending = week_dates[-1].strftime('%m/%d/%Y') if week_dates else ''
    project_names = ', '.join(sorted(repo_commits))
    # Fill fields in the first 4 paragraphs (assuming template structure)
    for para in doc.paragraphs:
        if 'FROM:' in para.text:
            for run in para.runs:
                if 'FROM:' in run.text:
                    run.text = f"FROM: Harsha Kocherla"
            if 'WEEK ENDING:' in para.text:
                para.text = para.text.split('WEEK ENDING:')[0] + f"WEEK ENDING: {week_ending}"
        elif 'WEEK ENDING:' in para.text:
            for run in para.runs:
                if 'WEEK ENDING:' in run.text:
                    run.text = f"WEEK ENDING: {week_ending}"
        elif 'CLIENT NAME:' in para.text:
            for run in para.runs:
                if 'CLIENT NAME:' in run.text:
                    run.text = f"CLIENT NAME: Best Buy"
            if 'PROJECT NAME:' in para.text:
                para.text = para.text.split('PROJECT NAME:')[0] + f"PROJECT NAME: {project_names}"
        elif 'PROJECT NAME:' in para.text:
            for run in para.runs:
                if 'PROJECT NAME:' in run.text:
                    run.text = f"PROJECT NAME: {project_names}"
    table = doc.tables[0]
    # Remove all rows except the header
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    border_settings = {
        "top":    {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "bottom": {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "left":   {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "right":  {"sz": 12, "val": "single", "color": "000000", "space": "0"},
    }
    for date, message in sorted(week_report.items()):
        row = table.add_row().cells
        row[0].text = date.strftime('%m/%d/%Y')  # DATE
        row[1].text = message                   # TASK
        for cell in row:
            set_cell_border(cell, **border_settings)
            set_cell_padding(cell)  # Add padding to each cell
    doc.save(output_path)

def main():
    if len(sys.argv) != 2:
        print("Usage: wsr <monday_date>")
        print("Example: wsr 06/02/2026")
        sys.exit(1)

    try:
        monday_date = parse(sys.argv[1]).date()
    except ValueError:
        print("Error: Invalid date format. Please use MM/DD/YYYY format.")
        sys.exit(1)

    if not is_monday(monday_date):
        print("Error: Please provide a Monday date.")
        sys.exit(1)

    workspace_path = os.path.expanduser("~/Workspaces")
    all_reports = process_workspace(workspace_path, monday_date)
    
    # For each week, collect repo names with commits for that week
    template_path = os.path.join(os.path.dirname(__file__), '../template/WSR_Template.docx')
    if os.path.exists(template_path):
        for week_monday, week_report in sorted(all_reports.items()):
            repo_commits = set()
            for item in Path(workspace_path).iterdir():
                if item.is_dir() and is_git_repo(item):
                    repo = git.Repo(item)
                    user_email, user_name = get_current_git_user(repo)
                    for date in week_report:
                        commits = get_commits_for_date(repo, date, user_email, user_name)
                        if commits:
                            repo_commits.add(item.name)
            output_docx = f"status_report_{week_monday.strftime('%Y%m%d')}.docx"
            fill_docx_report(template_path, output_docx, week_report, week_monday, repo_commits)
            print(f"Status report (docx) generated: {output_docx}")
    else:
        print(f"Template not found at {template_path}. Skipping docx generation.")

    # Optionally, still generate a single txt file for all weeks
    output_file = f"status_report_{monday_date.strftime('%Y%m%d')}.txt"
    generate_report(all_reports, output_file)
    print(f"\nStatus report generated: {output_file}")

if __name__ == "__main__":
    main() 